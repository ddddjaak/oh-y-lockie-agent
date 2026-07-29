"""
Apply fonts to pandoc-generated docx. Template provides all styling, header, footer.
This script only ensures correct EN/CN fonts on body text and tables.

Usage: python setup_docx.py input.docx [output.docx]
"""
import sys
from docx import Document
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

FONT_EN = "Arial"
FONT_CN = "微软雅黑"
FONT_CODE = "Cascadia Code"


def apply_fonts_to_document(doc):
    """Apply EN/CN fonts. Source Code paragraphs get Cascadia Code."""
    for paragraph in doc.paragraphs:
        is_code = paragraph.style and paragraph.style.name == 'Source Code'
        for run in paragraph.runs:
            if is_code:
                _force_run_fonts(run, FONT_CODE, FONT_CN)
            else:
                _set_run_fonts(run, FONT_EN, FONT_CN)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    is_code = paragraph.style and paragraph.style.name == 'Source Code'
                    for run in paragraph.runs:
                        if is_code:
                            _force_run_fonts(run, FONT_CODE, FONT_CN)
                        else:
                            _set_run_fonts(run, FONT_EN, FONT_CN)

    for section in doc.sections:
        for hf in [section.header, section.footer,
                   section.first_page_header, section.first_page_footer]:
            try:
                for p in hf.paragraphs:
                    for run in p.runs:
                        _set_run_fonts(run, FONT_EN, FONT_CN)
            except Exception:
                pass


def _set_run_fonts(run, font_en, font_cn):
    """Set font names on a run (fill-in, don't override existing)."""
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml('<w:rFonts %s/>' % nsdecls("w"))
        rPr.insert(0, rFonts)

    has_explicit = any(rFonts.get(qn(a)) for a in ['w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'])
    if has_explicit:
        return

    if not rFonts.get(qn('w:ascii')) and not rFonts.get(qn('w:asciiTheme')):
        rFonts.set(qn('w:ascii'), font_en)
    if not rFonts.get(qn('w:hAnsi')) and not rFonts.get(qn('w:hAnsiTheme')):
        rFonts.set(qn('w:hAnsi'), font_en)
    if not rFonts.get(qn('w:eastAsia')) and not rFonts.get(qn('w:eastAsiaTheme')):
        rFonts.set(qn('w:eastAsia'), font_cn)
    if not rFonts.get(qn('w:cs')) and not rFonts.get(qn('w:cstheme')):
        rFonts.set(qn('w:cs'), font_en)


def _force_run_fonts(run, font_en, font_cn):
    """Force-set font names, overriding theme/existing fonts (for code blocks)."""
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml('<w:rFonts %s/>' % nsdecls("w"))
        rPr.insert(0, rFonts)
    for attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme', 'w:cstheme']:
        if rFonts.get(qn(attr)):
            del rFonts.attrib[qn(attr)]
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:cs'), font_en)


def main():
    if len(sys.argv) < 2:
        print("Usage: python setup_docx.py input.docx [output.docx]")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else input_path

    print("Applying fonts (EN=%s, CN=%s, Code=%s)..." % (FONT_EN, FONT_CN, FONT_CODE))
    doc = Document(input_path)
    apply_fonts_to_document(doc)
    doc.save(output_path)
    print("Saved: %s" % output_path)


if __name__ == "__main__":
    main()
