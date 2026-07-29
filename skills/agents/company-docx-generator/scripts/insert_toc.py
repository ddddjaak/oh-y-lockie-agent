"""
Insert Word TOC field after '目录' or 'Table of Contents' heading.
Usage: python insert_toc.py input.docx
"""
import sys
import os
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

# TOC field code: \o "1-3" = outline levels 1-3, \h = hyperlinks, \z = hide tab leader, \u = use paragraph outline level
TOC_CODE = r' TOC \o "1-3" \h \z \u '
PLACEHOLDER_CN = '（打开文档后按 Ctrl+A 再按 F9 刷新目录）'
PLACEHOLDER_EN = '(Open the document, press Ctrl+A then F9 to update TOC)'
TOC_HEADINGS = ['目录', 'Table of Contents']


def insert_toc(doc_path):
    doc = Document(doc_path)

    # Find TOC heading
    toc_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() in TOC_HEADINGS and 'Heading' in p.style.name:
            toc_idx = i
            break

    if toc_idx is None:
        print('WARNING: No "目录" or "Table of Contents" heading found, skipping TOC insertion')
        return False

    # Create TOC paragraph after the heading
    h = doc.paragraphs[toc_idx]._element
    np = etree.SubElement(h.getparent(), qn('w:p'))
    h.addnext(np)

    placeholder = PLACEHOLDER_CN

    for item in [
        ('begin', None),
        ('instr', TOC_CODE),
        ('separate', None),
        ('text', placeholder),
        ('end', None),
    ]:
        r = etree.SubElement(np, qn('w:r'))
        if item[0] == 'begin':
            f = etree.SubElement(r, qn('w:fldChar'))
            f.set(qn('w:fldCharType'), 'begin')
        elif item[0] == 'instr':
            i = etree.SubElement(r, qn('w:instrText'))
            i.set(qn('xml:space'), 'preserve')
            i.text = item[1]
        elif item[0] == 'separate':
            f = etree.SubElement(r, qn('w:fldChar'))
            f.set(qn('w:fldCharType'), 'separate')
        elif item[0] == 'text':
            t = etree.SubElement(r, qn('w:t'))
            t.text = item[1]
            t.set(qn('xml:space'), 'preserve')
        elif item[0] == 'end':
            f = etree.SubElement(r, qn('w:fldChar'))
            f.set(qn('w:fldCharType'), 'end')

    print(f'TOC field inserted after P{toc_idx} ("{doc.paragraphs[toc_idx].text.strip()}" heading)')
    doc.save(doc_path)
    return True


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python insert_toc.py input.docx")
        sys.exit(1)
    insert_toc(sys.argv[1])
