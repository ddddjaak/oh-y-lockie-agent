"""
Post-process a pandoc-generated docx to add table borders, alternating row shading, and fonts.
Usage: python style_tables.py input.docx [output.docx]
"""
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from copy import deepcopy
import os

# Table styling
HEADER_BG = "0F2B46"      # Deep Navy header background
HEADER_FG = "FFFFFF"      # White header text
ODD_BG = "F1F5F9"         # Slate-50 for odd rows
EVEN_BG = "FFFFFF"        # White for even rows
BORDER_COLOR = "CBD5E1"   # Slate-200 borders
BORDER_SIZE = "4"         # 0.5pt in eighths of a point

# Font settings
FONT_EN = "Arial"             # English/numbers font
FONT_CN = "微软雅黑"           # Chinese font

def set_cell_shading(cell, color):
    """Set cell background shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = tcPr.find(qn('w:shd'))
    if shading is None:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
        tcPr.append(shading)
    else:
        shading.set(qn('w:fill'), color)
        shading.set(qn('w:val'), 'clear')

def set_cell_borders(cell, size=BORDER_SIZE, color=BORDER_COLOR):
    """Set cell borders on all four sides."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
            f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
            f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
            f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)
    else:
        # Update existing borders
        for edge in ['top', 'left', 'bottom', 'right']:
            edge_el = borders.find(qn(f'w:{edge}'))
            if edge_el is None:
                edge_el = parse_xml(
                    f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
                )
                borders.append(edge_el)
            else:
                edge_el.set(qn('w:val'), 'single')
                edge_el.set(qn('w:sz'), size)
                edge_el.set(qn('w:space'), '0')
                edge_el.set(qn('w:color'), color)

def set_table_borders(table, size=BORDER_SIZE, color=BORDER_COLOR):
    """Set outer table borders."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

    borders = tblPr.find(qn('w:tblBorders'))
    if borders is None:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
            f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
            f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
            f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
            f'  <w:insideH w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
            f'  <w:insideV w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)

def set_run_color(run, color):
    """Set the text color of a run."""
    rPr = run._r.get_or_add_rPr()
    color_el = rPr.find(qn('w:color'))
    if color_el is None:
        color_el = parse_xml(f'<w:color {nsdecls("w")} w:val="{color}"/>')
        rPr.append(color_el)
    else:
        color_el.set(qn('w:val'), color)

def set_run_font(run, font_en=FONT_EN, font_cn=FONT_CN):
    """Set font for a run - English font for ascii, Chinese font for eastAsia."""
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)
    rFonts.set(qn('w:eastAsia'), font_cn)

def set_document_fonts(doc, font_en=FONT_EN, font_cn=FONT_CN):
    """Set fonts for the entire document (all paragraphs and tables)."""
    # Set default font in styles
    styles_elem = doc.styles.element
    doc_defaults = styles_elem.find(qn('w:docDefaults'))
    if doc_defaults is not None:
        rPr_default = doc_defaults.find(qn('w:rPrDefault'))
        if rPr_default is None:
            rPr_default = parse_xml(f'<w:rPrDefault {nsdecls("w")}/>')
            doc_defaults.append(rPr_default)
        rPr = rPr_default.find(qn('w:rPr'))
        if rPr is None:
            rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
            rPr_default.append(rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), font_en)
        rFonts.set(qn('w:hAnsi'), font_en)
        rFonts.set(qn('w:eastAsia'), font_cn)

    # Apply font to all paragraphs in document body (skip Source Code and cover title)
    for paragraph in doc.paragraphs:
        if paragraph.style and paragraph.style.name == 'Source Code':
            continue
        for run in paragraph.runs:
            # Skip runs that already have explicit fonts set (e.g. cover title)
            rPr = run._r.find(qn('w:rPr'))
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    has_explicit = any(rFonts.get(qn(a)) for a in ['w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'])
                    if has_explicit:
                        continue
            set_run_font(run, font_en, font_cn)

    # Apply font to all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.style and paragraph.style.name == 'Source Code':
                        continue
                    for run in paragraph.runs:
                        set_run_font(run, font_en, font_cn)

def enforce_cell_wrap(table):
    """Ensure all cells have word-wrap enabled and reasonable font size.

    Wide tables (6+ columns): reduce font to 8pt to prevent overflow.
    Narrow columns: enable CJK line breaking so Chinese text wraps properly.
    """
    num_cols = len(table.columns)
    if num_cols == 0:
        return

    # Determine font size based on column count
    if num_cols >= 8:
        cell_font_size = 15  # 7.5pt in half-points
    elif num_cols >= 6:
        cell_font_size = 16  # 8pt
    elif num_cols >= 4:
        cell_font_size = 18  # 9pt
    else:
        cell_font_size = 20  # 10pt

    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()

            # Enable word wrap (disable noWrap)
            noWrap = tcPr.find(qn('w:noWrap'))
            if noWrap is not None:
                tcPr.remove(noWrap)

            # Set cell width type to auto for flexibility
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is not None:
                tcW.set(qn('w:type'), 'auto')

            # Apply font size to all paragraphs in cell
            for para in cell.paragraphs:
                for run in para.runs:
                    rPr = run._r.get_or_add_rPr()
                    sz = rPr.find(qn('w:sz'))
                    if sz is None:
                        sz = parse_xml('<w:sz %s w:val="%d"/>' % (nsdecls("w"), cell_font_size))
                        rPr.append(sz)
                    else:
                        current_sz = int(sz.get(qn('w:val')))
                        if current_sz > cell_font_size:
                            sz.set(qn('w:val'), str(cell_font_size))


def enforce_page_width(table):
    """Clamp table width to 100% page width. Prevent horizontal overflow."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml('<w:tblPr %s/>' % nsdecls("w"))
        tbl.insert(0, tblPr)

    # Set table width to 100% (5000 pct)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = parse_xml('<w:tblW %s w:w="5000" w:type="pct"/>' % nsdecls("w"))
        tblPr.append(tblW)
    else:
        tblW.set(qn('w:w'), '5000')
        tblW.set(qn('w:type'), 'pct')

    # Remove any preferred table width that might override
    for attr in ['w:w', 'w:type']:
        existing = tblPr.find(qn(attr))
        # Keep our values

    # Add table indent of 0 to prevent right-side overflow from negative indents
    tblInd = tblPr.find(qn('w:tblInd'))
    if tblInd is not None:
        tblInd.set(qn('w:w'), '0')

    # Ensure table layout is autofit (not fixed)
    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is None:
        tblLayout = parse_xml('<w:tblLayout %s w:type="autofit"/>' % nsdecls("w"))
        tblPr.append(tblLayout)
    else:
        tblLayout.set(qn('w:type'), 'autofit')


def style_tables(doc):
    """Style all tables in the document with borders, alternating row colors, and auto-fit."""
    for table in doc.tables:
        num_rows = len(table.rows)
        if num_rows == 0:
            continue

        # Auto-fit: enforce page width, enable wrapping, scale fonts for wide tables
        enforce_page_width(table)
        enforce_cell_wrap(table)

        # Set table-level borders (inside lines)
        set_table_borders(table)

        # Determine if first row is a header
        # Pandoc typically marks header rows by making text bold
        # We identify header rows by checking if the first row has bold text
        # or if the table has a special style

        has_header = False
        if num_rows >= 1:
            first_row = table.rows[0]

            # Method 1: Check OOXML tblHeader marker (pandoc sets this for markdown table headers)
            trPr = first_row._tr.find(qn('w:trPr'))
            if trPr is not None and trPr.find(qn('w:tblHeader')) is not None:
                has_header = True

            # Method 2: Check if >50% of cells have bold text (fallback)
            if not has_header:
                bold_count = 0
                total_cells = len(first_row.cells)
                for cell in first_row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.bold:
                                bold_count += 1
                                break
                if total_cells > 0 and bold_count / total_cells > 0.5:
                    has_header = True

        for row_idx, row in enumerate(table.rows):
            is_header = has_header and row_idx == 0
            is_odd = row_idx % 2 == 1

            # Determine row shading
            if is_header:
                bg = HEADER_BG
                fg = HEADER_FG
            elif is_odd:
                bg = ODD_BG
                fg = None  # keep default text color
            else:
                bg = EVEN_BG
                fg = None

            for cell in row.cells:
                # Set cell borders
                set_cell_borders(cell)

                # Set cell shading
                if bg:
                    set_cell_shading(cell, bg)

                # Set text color for header
                if fg:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            set_run_color(run, fg)

def main():
    if len(sys.argv) < 2:
        print("Usage: python style_tables.py input.docx [output.docx]")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else input_path

    print(f"Processing: {input_path}")
    doc = Document(input_path)

    print(f"Setting fonts: EN={FONT_EN}, CN={FONT_CN}")
    set_document_fonts(doc)

    print(f"Found {len(doc.tables)} tables")
    style_tables(doc)

    doc.save(output_path)
    print(f"Saved: {output_path}")

if __name__ == '__main__':
    main()
