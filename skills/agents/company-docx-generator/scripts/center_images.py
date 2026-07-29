"""
Center images, image captions, and table captions in a docx.
Scale oversized images to fit within page margins.

Image captions: the paragraph right after each image (pandoc renders alt text here) → centered
Table captions: the paragraph right before each table that looks like a caption → centered

Usage: python center_images.py input.docx
"""
import sys
from docx import Document
from docx.shared import Inches, Emu, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Maximum image dimensions (A4 portrait, 2.5cm margins on each side)
MAX_IMAGE_WIDTH_EMU = int(Cm(16.0))   # ~6.3 inches
MAX_IMAGE_HEIGHT_EMU = int(Cm(22.0))  # ~8.7 inches — leave room for caption


def constrain_image_size(doc):
    """Scale down images that exceed max width/height, preserving aspect ratio."""
    oversized = 0

    # Find all inline images in drawings
    for para in doc.paragraphs:
        for run in para._element.iterchildren(qn('w:r')):
            drawing = run.find(qn('w:drawing'))
            if drawing is None:
                continue

            # Navigate: drawing → inline → extent (cx, cy)
            inline = drawing.find(qn('wp:inline'))
            if inline is None:
                # Could be anchor (floating image) — skip for now
                continue

            extent = inline.find(qn('wp:extent'))
            if extent is None:
                continue

            cx = int(extent.get('cx', '0'))
            cy = int(extent.get('cy', '0'))

            if cx <= 0 or cy <= 0:
                continue

            # Check if oversized
            if cx <= MAX_IMAGE_WIDTH_EMU and cy <= MAX_IMAGE_HEIGHT_EMU:
                continue

            # Scale down, preserving aspect ratio
            scale_w = MAX_IMAGE_WIDTH_EMU / cx
            scale_h = MAX_IMAGE_HEIGHT_EMU / cy
            scale = min(scale_w, scale_h, 1.0)

            new_cx = int(cx * scale)
            new_cy = int(cy * scale)
            extent.set('cx', str(new_cx))
            extent.set('cy', str(new_cy))

            # Also scale any transform in the graphic
            graphic = inline.find(qn('a:graphic'))
            if graphic is not None:
                for xfrm in graphic.iter(qn('a:xfrm')):
                    ext_elem = xfrm.find(qn('a:ext'))
                    if ext_elem is not None:
                        ext_elem.set('cx', str(new_cx))
                        ext_elem.set('cy', str(new_cy))

            oversized += 1

    if oversized:
        print(f'  Scaled {oversized} oversized image(s) to fit page margins')
    return oversized


def center_figures(doc_path):
    doc = Document(doc_path)

    # Step 0: Constrain oversized images before centering
    constrain_image_size(doc)

    # Step 1: Find all image paragraphs and the captions after them
    img_indices = []
    for i, para in enumerate(doc.paragraphs):
        drawings = para._element.findall('.//' + qn('w:drawing'))
        if drawings:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_indices.append(i)

    # Step 2: Center the caption paragraph right after each image
    # Pandoc renders ![alt text](path) as: [image paragraph] + [alt text paragraph]
    captions_centered = 0
    for idx in img_indices:
        next_idx = idx + 1
        if next_idx < len(doc.paragraphs):
            next_para = doc.paragraphs[next_idx]
            text = next_para.text.strip()
            # Pandoc caption paragraph: non-empty, not a heading, contains image alt text
            # Typically starts with "图" or "Fig" or contains the alt text
            if text and 'Heading' not in next_para.style.name:
                # Check it's not another image
                if not next_para._element.findall('.//' + qn('w:drawing')):
                    next_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    captions_centered += 1

    # Step 3: Find tables and center caption paragraphs above them
    # A table caption is identified by starting with "表" or "Table"
    table_captions = 0
    CAPTION_KEYWORDS = ['表', 'Table ']

    for table in doc.tables:
        tbl_elem = table._element
        body = tbl_elem.getparent()
        tbl_index = list(body).index(tbl_elem)

        # Look backwards from table position to find the nearest paragraph
        for offset in range(1, min(tbl_index + 1, 4)):  # Only look up to 3 elements back
            prev_elem = body[tbl_index - offset]
            if prev_elem.tag == qn('w:p'):
                for para in doc.paragraphs:
                    if para._element is prev_elem:
                        text = para.text.strip()
                        # Only center if it looks like a table caption
                        if text and 'Heading' not in para.style.name:
                            if any(text.startswith(kw) for kw in CAPTION_KEYWORDS):
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                table_captions += 1
                        break
                break  # Only consider the paragraph immediately before the table

    doc.save(doc_path)
    print(f'Centered {len(img_indices)} image(s) + {captions_centered} image caption(s) + {table_captions} table caption(s)')
    return len(img_indices)


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python center_images.py input.docx")
        sys.exit(1)
    center_figures(sys.argv[1])
