"""
Extract structure info from a company template docx.
Uses the docx skill's unpack approach for raw XML access,
falling back to python-docx for quick analysis.

Usage: python extract_template.py <template.docx>
"""
import sys
import os
import zipfile
import re
from docx import Document
from docx.oxml.ns import qn

TEMPLATE_PATH = None  # set from CLI arg


def analyze_via_docx(path):
    """Quick analysis via python-docx — section, styles, header/footer overview."""
    doc = Document(path)
    print("=== SECTIONS ===")
    for i, s in enumerate(doc.sections):
        pw = int(s.page_width / 914400 * 100) / 100
        ph = int(s.page_height / 914400 * 100) / 100
        ml = int(s.left_margin / 914400 * 100) / 100
        mr = int(s.right_margin / 914400 * 100) / 100
        mt = int(s.top_margin / 914400 * 100) / 100
        mb = int(s.bottom_margin / 914400 * 100) / 100
        print("  Page: %.1f\" x %.1f\"  Margins: L=%.2f R=%.2f T=%.2f B=%.2f" % (pw, ph, ml, mr, mt, mb))
        print("  diff_first_page=%s" % s.different_first_page_header_footer)

        h = s.header
        print("  Header (%d paras):" % len(h.paragraphs))
        for j, p in enumerate(h.paragraphs):
            texts = [r.text[:40] for r in p.runs if r.text.strip()]
            has_img = any(r._r.findall('.//' + qn('a:blip')) for r in p.runs)
            print("    P%d: img=%s texts=%s" % (j, has_img, texts[:2]))

        f = s.footer
        print("  Footer (%d paras):" % len(f.paragraphs))
        for j, p in enumerate(f.paragraphs):
            texts = [r.text[:40] for r in p.runs if r.text.strip()]
            print("    P%d: %s" % (j, texts[:3]))

    # Heading styles
    print()
    print("=== HEADING STYLES ===")
    for style in doc.styles:
        if style.type is not None and 'PARAGRAPH' in str(style.type):
            name = (style.name or '').lower()
            if 'heading' in name:
                font = style.font
                info = ['id=%s' % style.style_id]
                if font:
                    if font.size: info.append('size=%spt' % (int(font.size) / 20.0))
                    if font.name: info.append('font=%s' % font.name)
                    if font.bold: info.append('bold')
                # Check outline level via raw XML
                pPr = style._element.find(qn('w:pPr'))
                if pPr is not None:
                    ol = pPr.find(qn('w:outlineLvl'))
                    if ol is not None:
                        info.append('level=%s' % ol.get(qn('w:val')))
                print("  %-20s %s" % (name, ', '.join(info)))

    # Numbering
    print()
    print("=== NUMBERING ===")
    try:
        num_part = doc.part.numbering_part
        print("  present=%s" % (num_part is not None))
    except:
        print("  not accessible")


def analyze_via_unpack(path):
    """Deep analysis via ZIP unpack — raw XML for header/footer details."""
    print("=== RAW HEADER/FOOTER (via unpack) ===")
    with zipfile.ZipFile(path) as z:
        for name in sorted(z.namelist()):
            if 'header' in name.lower() and name.endswith('.rels'):
                xml = z.read(name).decode('utf-8')
                imgs = re.findall(r'Target="([^"]*)"', xml)
                if imgs:
                    print("  %s -> images: %s" % (name, imgs))
            if 'header' in name.lower() and not name.endswith('.rels') and name.endswith('.xml'):
                xml = z.read(name).decode('utf-8')
                # Extract key info
                has_drawing = 'wp:drawing' in xml or 'w:drawing' in xml
                has_pict = 'w:pict' in xml or 'v:shape' in xml
                has_logo = 'r:embed=' in xml
                texts = re.findall(r'<w:t[^>]*>(.*?)</w:t>', xml)
                text_summary = ''.join(t for t in texts if t.strip())[:80]
                print("  %s: drawing=%s pict=%s logo=%s text=\"%s\"" % (
                    name, has_drawing, has_pict, has_logo, text_summary))
            if 'footer' in name.lower() and not name.endswith('.rels') and name.endswith('.xml'):
                xml = z.read(name).decode('utf-8')
                texts = re.findall(r'<w:t[^>]*>(.*?)</w:t>', xml)
                text_summary = ''.join(t for t in texts if t.strip())[:80]
                # Find doc number patterns
                doc_nums = re.findall(r'FM-RD-\d+\s+A\d+', ''.join(texts))
                print("  %s: text=\"%s\" doc#=%s" % (
                    name, text_summary, doc_nums))
        # Media
        media = [n for n in z.namelist() if 'media' in n]
        if media:
            print("  Media files: %d images" % len(media))


def main():
    global TEMPLATE_PATH
    if len(sys.argv) < 2:
        print("Usage: python extract_template.py <template.docx>")
        print("Analyzes a company template docx — sections, styles, headers, footers.")
        sys.exit(1)

    TEMPLATE_PATH = sys.argv[1]
    if not os.path.exists(TEMPLATE_PATH):
        print("File not found: %s" % TEMPLATE_PATH)
        sys.exit(1)

    print("Template: %s" % TEMPLATE_PATH)
    print()

    analyze_via_docx(TEMPLATE_PATH)
    print()
    analyze_via_unpack(TEMPLATE_PATH)


if __name__ == "__main__":
    main()
