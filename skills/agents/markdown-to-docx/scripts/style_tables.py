"""
Post-process a pandoc-generated docx to add table borders and alternating row shading.
Usage: python style_tables.py input.docx [output.docx]
"""
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from copy import deepcopy
import os

HEADER_BG = "2F5496"      # Dark blue header background
HEADER_FG = "FFFFFF"      # White header text
ODD_BG = "D6E4F0"         # Light blue for odd rows
EVEN_BG = "FFFFFF"        # White for even rows
BORDER_COLOR = "808080"   # Gray borders
BORDER_SIZE = "4"         # 0.5pt in eighths of a point

def set_cell_shading(cell, color):
    """Set cell background shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = tcPr.find(qn('w:shd'))
    if shading is None:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
        tcPr.append(shading)
    else:
        shading.set(qn('w:fill'), color)
        shading.set(qn('w:val'), 'clear')

def set_cell_borders(cell, size=BORDER_SIZE, color=BORDER_COLOR):
    """Set cell borders on all four sides."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
            f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
            f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
            f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)
    else:
        # Update existing borders
        for edge in ['top', 'left', 'bottom', 'right']:
            edge_el = borders.find(qn(f'w:{edge}'))
            if edge_el is None:
                edge_el = parse_xml(
                    f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
                )
                borders.append(edge_el)
            else:
                edge_el.set(qn('w:val'), 'single')
                edge_el.set(qn('w:sz'), size)
                edge_el.set(qn('w:space'), '0')
                edge_el.set(qn('w:color'), color)

def set_table_borders(table, size=BORDER_SIZE, color=BORDER_COLOR):
    """Set outer table borders."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

    borders = tblPr.find(qn('w:tblBorders'))
    if borders is None:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
            f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
            f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
            f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
            f'  <w:insideH w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
            f'  <w:insideV w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)

def set_run_color(run, color):
    """Set the text color of a run."""
    rPr = run._r.get_or_add_rPr()
    color_el = rPr.find(qn('w:color'))
    if color_el is None:
        color_el = parse_xml(f'<w:color {nsdecls("w")} w:val="{color}"/>')
        rPr.append(color_el)
    else:
        color_el.set(qn('w:val'), color)

def set_paragraph_style_in_cell(cell, font_name='Microsoft YaHei', font_size=Pt(9)):
    """Ensure cell text uses specified font."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = font_size
            run.font.name = font_name

def style_tables(doc):
    """Style all tables in the document with borders and alternating row colors."""
    for table in doc.tables:
        num_rows = len(table.rows)
        if num_rows == 0:
            continue

        # Set table-level borders (inside lines)
        set_table_borders(table)

        # Determine if first row is a header
        # Pandoc typically marks header rows by making text bold
        # We identify header rows by checking if the first row has bold text
        # or if the table has a special style

        has_header = False
        if num_rows >= 1:
            first_row = table.rows[0]

            # Method 1: Check OOXML tblHeader marker (pandoc sets this for markdown table headers)
            trPr = first_row._tr.find(qn('w:trPr'))
            if trPr is not None and trPr.find(qn('w:tblHeader')) is not None:
                has_header = True

            # Method 2: Check if >50% of cells have bold text (fallback)
            if not has_header:
                bold_count = 0
                total_cells = len(first_row.cells)
                for cell in first_row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.bold:
                                bold_count += 1
                                break
                if total_cells > 0 and bold_count / total_cells > 0.5:
                    has_header = True

        for row_idx, row in enumerate(table.rows):
            is_header = has_header and row_idx == 0
            is_odd = row_idx % 2 == 1

            # Determine row shading
            if is_header:
                bg = HEADER_BG
                fg = HEADER_FG
            elif is_odd:
                bg = ODD_BG
                fg = None  # keep default text color
            else:
                bg = EVEN_BG
                fg = None

            for cell in row.cells:
                # Set cell borders
                set_cell_borders(cell)

                # Set cell shading
                if bg:
                    set_cell_shading(cell, bg)

                # Set text color for header
                if fg:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            set_run_color(run, fg)

def main():
    if len(sys.argv) < 2:
        print("Usage: python style_tables.py input.docx [output.docx]")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else input_path

    print(f"Processing: {input_path}")
    doc = Document(input_path)

    print(f"Found {len(doc.tables)} tables")
    style_tables(doc)

    doc.save(output_path)
    print(f"Saved: {output_path}")

if __name__ == '__main__':
    main()
